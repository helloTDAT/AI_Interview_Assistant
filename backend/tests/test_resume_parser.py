from pathlib import Path

from app.models import ResumeFileKind
from app.services.resume_parser import ResumeParser


def test_image_resume_uses_ocr_fallback(tmp_path: Path):
    image = tmp_path / "resume.png"
    image.write_bytes(b"not really an image but enough for extension routing")

    profile = ResumeParser().parse(image)

    assert profile.source_kind == ResumeFileKind.image
    assert profile.parse_confidence < 0.6
    assert profile.warnings
    assert "OCR" in profile.raw_text


def test_docx_resume_extracts_text(tmp_path: Path):
    from docx import Document

    docx = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("求职意向：后端开发工程师")
    doc.add_paragraph("技能：Python MySQL Redis FastAPI")
    doc.save(str(docx))

    profile = ResumeParser().parse(docx)

    assert profile.source_kind == ResumeFileKind.docx
    assert profile.candidate_name == "张三"
    assert "Python" in profile.skills
    assert profile.target_position == "后端开发工程师"


def test_docx_project_headings_are_not_extracted_as_projects(tmp_path: Path):
    from docx import Document

    docx = tmp_path / "project_headings.docx"
    doc = Document()
    doc.add_paragraph("李四")
    doc.add_paragraph("求职意向：后端开发工程师")
    doc.add_paragraph("项目经历")
    doc.add_paragraph("【项目背景】")
    doc.add_paragraph("优化")
    doc.add_paragraph("提升")
    doc.add_paragraph("项目名称：分布式 IM 后端系统")
    doc.add_paragraph("负责使用 Netty 和 Redis Cluster 实现消息路由，支持 10 万连接压测。")
    doc.save(str(docx))

    profile = ResumeParser().parse(docx)

    joined_projects = "\n".join(profile.projects)
    assert "分布式 IM 后端系统" in joined_projects
    assert "【项目背景】" not in joined_projects
    assert "优化\n提升" not in joined_projects
