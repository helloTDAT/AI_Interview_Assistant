from pathlib import Path

from fastapi.testclient import TestClient

from app.api import routes
from app.core.database import review_db
from app.core.storage import store
from app.main import app
from app.services.ai_clients import LLMCallResult, LLMClient


client = TestClient(app)


def test_health():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json()["status"] == "ok"


def test_upload_image_resume(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(routes.resume_agent, "llm", LLMClient(api_key=""))
    image = tmp_path / "resume.jpg"
    image.write_bytes(b"fake image")

    with image.open("rb") as handle:
        response = client.post(
            "/files/resume",
            files={"file": ("resume.jpg", handle, "image/jpeg")},
            data={
                "user_id": "api-user",
                "target_position": "AI算法工程师",
                "user_instruction": "请重点优化项目经历。",
            },
        )

    assert response.status_code == 200
    body = response.json()
    assert body["profile"]["source_kind"] == "image"
    assert body["needs_user_confirmation"] is True
    assert body["job_fit"]["target_position"] == "AI算法工程师"
    assert "机器学习" in body["job_fit"]["expected_skills"]
    assert any("项目经历" in item or "项目" in item for item in body["recommendations"])


def test_mock_interview_flow():
    start = client.post("/interviews/mock/start", json={"user_id": "api-user", "target_position": "后端开发工程师"})
    assert start.status_code == 200
    session_id = start.json()["session_id"]

    answer = client.post(
        "/interviews/mock/message",
        json={"session_id": session_id, "answer": "我负责后端接口设计，使用 FastAPI 和 MySQL 完成核心模块。"},
    )
    assert answer.status_code == 200
    assert answer.json()["feedback"]


def test_mock_interview_modes_and_probing_fields():
    latest_start = None
    for mode in ["tech_blitz", "project_deep_dive", "behavioral"]:
        latest_start = client.post(
            "/interviews/mock/start",
            json={"user_id": "api-user", "target_position": "AI算法工程师", "mode": mode},
        )
        assert latest_start.status_code == 200
        body = latest_start.json()
        assert body["mode"] == mode
        assert body["state"] == "Greeting"
        assert "pressure_level" in body
        assert body["skill_scores"]

    answer = client.post(
        "/interviews/mock/message",
        json={"session_id": latest_start.json()["session_id"], "answer": "我用了 PyTorch 做模型训练。"},
    )
    assert answer.status_code == 200
    payload = answer.json()
    assert payload["state"] in {"Probing", "NewQuestion", "ReverseQA", "Closing"}
    assert payload["probing_reason"]
    assert "PyTorch" in payload["detected_keywords"]
    assert payload["answer_depth_score"] >= 0


def test_mock_interview_without_resume_uses_generic_question_not_topic_title(monkeypatch):
    class FakeQuestionAgent:
        def generate_for_resume(self, profile, target_position, count=5):
            raise AssertionError("question bank should not be used before resume upload")

    original_agent = routes.mock_agent.question_agent
    routes.mock_agent.question_agent = FakeQuestionAgent()
    try:
        start = client.post(
            "/interviews/mock/start",
            json={"user_id": "topic-user", "target_position": "AI算法工程师", "mode": "project_deep_dive"},
        )
    finally:
        routes.mock_agent.question_agent = original_agent

    assert start.status_code == 200
    question = start.json()["question"]
    assert "从关系型数据库构建时间窗口特征" not in question
    assert "Agent 智能体" in question
    assert start.json()["question_type"] == "role_core"


def test_mock_interview_with_resume_targets_resume_project():
    store.resume_reports["resume-mock-user"] = {
        "profile": {
            "projects": ["智能医疗诊断平台：负责 Mamba 模型接入和 Vue 前端渲染优化"],
            "skills": ["Vue", "Mamba", "Python"],
            "raw_text": "参与智能医疗诊断平台，负责模型接入和前端渲染。",
        }
    }
    start = client.post(
        "/interviews/mock/start",
        json={"user_id": "resume-mock-user", "target_position": "AI算法工程师", "mode": "project_deep_dive"},
    )

    assert start.status_code == 200
    question = start.json()["question"]
    assert start.json()["question_type"] == "project"
    assert "智能医疗诊断平台" in question
    assert "职责" in question
    assert "应聘AI算法工程师" in question


def test_mock_interview_with_resume_covers_multiple_dimensions():
    store.resume_reports["coverage-user"] = {
        "profile": {
            "projects": ["智能医疗诊断平台：负责 Mamba 模型接入和 Vue 前端渲染优化"],
            "skills": ["Vue", "Mamba", "Python", "RAG"],
            "raw_text": "参与智能医疗诊断平台，负责模型接入和前端渲染。",
        }
    }
    start = client.post(
        "/interviews/mock/start",
        json={"user_id": "coverage-user", "target_position": "AI算法工程师", "mode": "project_deep_dive"},
    ).json()
    session_id = start["session_id"]
    seen_types = {start["question_type"]}
    seen_phases = {start.get("phase_label")}
    for index in range(7):
        response = client.post(
            "/interviews/mock/message",
            json={
                "session_id": session_id,
                "answer": (
                    f"第 {index} 轮回答：我先说明背景和目标，再讲方案、关键步骤、难点、风险、验证指标，"
                    "包括 Agent、RAG、向量、索引、复杂度、边界情况和 30% 效果提升。"
                ),
            },
        ).json()
        seen_types.add(response["question_type"])
        seen_phases.add(response.get("phase_label"))

    assert "project" in seen_types
    assert "foundation" in seen_types
    assert "algorithm" in seen_types
    assert {"项目暖场", "项目基础", "项目技术基础", "项目实现细节", "项目深挖"}.issubset(seen_phases)


def test_mock_interview_reaches_reverse_qa_and_summary():
    start = client.post(
        "/interviews/mock/start",
        json={"user_id": "api-user", "target_position": "后端开发工程师", "mode": "project_deep_dive"},
    )
    session_id = start.json()["session_id"]
    latest = None
    for index in range(8):
        latest = client.post(
            "/interviews/mock/message",
            json={
                "session_id": session_id,
                "answer": f"第 {index} 轮回答：我说明背景、方案、难点、结果和 30% 性能提升，并解释 Redis、MySQL 与队列取舍。",
            },
        ).json()
        if latest["state"] == "ReverseQA":
            break
    assert latest["state"] == "ReverseQA"
    assert latest["reverse_question_prompt"]

    close = client.post(
        "/interviews/mock/message",
        json={"session_id": session_id, "answer": "我想了解团队的技术栈、业务目标和新人培养方式。"},
    )
    assert close.status_code == 200
    close_body = close.json()
    assert close_body["finished"] is True
    assert close_body["final_report"]["dimension_scores"]


def test_mock_interview_finish_interrupts_and_keeps_session_closed():
    start = client.post(
        "/interviews/mock/start",
        json={"user_id": "finish-user", "target_position": "后端开发工程师", "mode": "project_deep_dive"},
    )
    session_id = start.json()["session_id"]

    finish = client.post(
        "/interviews/mock/finish",
        json={"session_id": session_id, "reason": "user_requested"},
    )
    assert finish.status_code == 200
    body = finish.json()
    assert body["finished"] is True
    assert body["state"] == "Closing"
    assert body["question_type"] == "summary"
    assert "提前结束" in body["final_report"]["summary"]
    assert body["final_report"]["dimension_scores"]

    after_finish = client.post(
        "/interviews/mock/message",
        json={"session_id": session_id, "answer": "我继续回答也不应该触发追问。"},
    )
    assert after_finish.status_code == 200
    after_body = after_finish.json()
    assert after_body["finished"] is True
    assert after_body["state"] == "Closing"
    assert after_body["question"] == body["question"]


def test_mock_interview_finish_missing_session_returns_404():
    response = client.post(
        "/interviews/mock/finish",
        json={"session_id": "missing-session", "reason": "user_requested"},
    )
    assert response.status_code == 404


def test_upload_resume_accepts_jd_and_llm_fields(tmp_path: Path, monkeypatch):
    class FakeLLM:
        def analyze_resume_report(self, profile, target_position, user_instruction=None, job_description=None):
            assert "PyTorch" in job_description
            return LLMCallResult(
                ok=True,
                data={
                    "quality_score": 90,
                    "dimension_scores": {"专业技能": 88, "岗位匹配": 85},
                    "job_fit": {
                        "target_position": target_position,
                        "fit_score": 85,
                        "expected_skills": ["Python", "PyTorch"],
                        "matched_skills": ["Python"],
                        "missing_skills": ["PyTorch"],
                    },
                    "recommendations": ["围绕 JD 补充 PyTorch 项目证据。"],
                    "jd_diagnosis": {
                        "enabled": True,
                        "match_rate": 85,
                        "core_requirements": ["Python", "PyTorch"],
                        "matched_items": ["Python"],
                        "missing_items": ["PyTorch"],
                        "suggestions": ["补充模型训练细节。"],
                    },
                    "interview_risks": [{"question": "请解释模型训练细节。"}],
                    "logic_gaps": [],
                    "reading_experience": {"signal_to_noise_score": 80},
                    "star_optimizations": [{"before": "做项目", "after": "建议按 STAR 改写", "action_note": "补充结果"}],
                },
            )

    monkeypatch.setattr(routes.resume_agent, "llm", FakeLLM())
    resume = tmp_path / "resume.docx"
    from docx import Document

    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("技能：Python SQL")
    doc.add_paragraph("项目经历：参与算法项目。")
    doc.save(str(resume))

    with resume.open("rb") as handle:
        response = client.post(
            "/files/resume",
            files={"file": ("resume.docx", handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={
                "user_id": "api-user",
                "target_position": "AI算法工程师",
                "user_instruction": "请重点优化项目经历",
                "job_description": "要求 Python、PyTorch、模型评估经验",
            },
        )

    assert response.status_code == 200
    body = response.json()
    assert body["analysis_engine"] == "llm"
    assert body["jd_diagnosis"]["missing_items"] == ["PyTorch"]
    assert body["interview_risks"][0]["question"]
    assert body["star_optimizations"][0]["after"]


def test_current_resume_and_cross_agent_reuse_after_upload(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(routes.resume_agent, "llm", LLMClient(api_key=""))
    resume = tmp_path / "resume.docx"
    from docx import Document

    doc = Document()
    doc.add_paragraph("李四")
    doc.add_paragraph("专业技能：Python Redis MySQL")
    doc.add_paragraph("项目经历1")
    doc.add_paragraph("智能面试助手平台：负责后端接口、RAG 检索和练习题推荐模块，实现个性化训练流程。")
    doc.save(str(resume))

    with resume.open("rb") as handle:
        upload = client.post(
            "/files/resume",
            files={"file": ("resume.docx", handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"user_id": "shared-resume-user", "target_position": "后端开发工程师", "user_instruction": "请分析简历"},
        )

    assert upload.status_code == 200
    current = client.get("/resume/current/shared-resume-user")
    assert current.status_code == 200
    current_body = current.json()["resume"]
    assert current_body["available"] is True
    assert current_body["target_position"] == "后端开发工程师"

    mock = client.post(
        "/interviews/mock/start",
        json={"user_id": "shared-resume-user", "target_position": "后端开发工程师", "mode": "project_deep_dive"},
    )
    assert mock.status_code == 200
    mock_body = mock.json()
    assert mock_body["resume_context_used"] is True
    assert mock_body["resume_context_summary"]
    assert "未读取" not in mock_body["resume_context_summary"]

    feed = client.post(
        "/learning/feed",
        json={"user_id": "shared-resume-user", "target_position": "后端开发工程师", "limit": 4},
    )
    assert feed.status_code == 200
    feed_body = feed.json()
    assert feed_body["profile_used"] is True
    assert "Python" in feed_body["profile_summary"]["skills"]

    second = tmp_path / "resume_v2.docx"
    doc2 = Document()
    doc2.add_paragraph("李四")
    doc2.add_paragraph("专业技能：Java Spring Redis")
    doc2.add_paragraph("项目经历1")
    doc2.add_paragraph("订单系统：负责 Java 接口和 Spring 服务拆分。")
    doc2.save(str(second))
    with second.open("rb") as handle:
        overwrite = client.post(
            "/files/resume",
            files={"file": ("resume_v2.docx", handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"user_id": "shared-resume-user", "target_position": "Java 开发工程师", "user_instruction": "更新简历画像"},
        )
    assert overwrite.status_code == 200
    updated = client.get("/resume/current/shared-resume-user").json()["resume"]
    assert updated["target_position"] == "Java 开发工程师"
    assert "Java" in updated["skills"]


def test_mock_interview_uses_raw_text_or_skills_when_projects_are_empty():
    store.resume_reports["raw-only-user"] = {
        "profile": {
            "projects": [],
            "internships": [],
            "skills": ["Python", "Redis", "MySQL"],
            "raw_text": "参与智能面试助手平台，负责 Python 后端接口、Redis 缓存和 MySQL 数据建模，实现练习推荐流程。",
            "parse_confidence": 0.86,
        },
        "job_fit": {"target_position": "后端开发工程师"},
        "quality_score": 82,
    }

    response = client.post(
        "/interviews/mock/start",
        json={"user_id": "raw-only-user", "target_position": "后端开发工程师", "mode": "project_deep_dive"},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["resume_context_used"] is True
    assert body["anchor_project"]
    assert "Python" in body["resume_context_summary"]


def test_current_resume_reports_missing_when_no_active_profile():
    store.resume_reports.pop("missing-resume-user", None)

    response = client.get("/resume/current/missing-resume-user")

    assert response.status_code == 200
    assert response.json()["resume"]["available"] is False


def test_audio_analysis_task(tmp_path: Path):
    audio = tmp_path / "interview.wav"
    audio.write_bytes(b"fake wav")

    with audio.open("rb") as handle:
        response = client.post(
            "/interviews/audio",
            files={"file": ("interview.wav", handle, "audio/wav")},
            data={"user_id": "audio-user"},
        )

    assert response.status_code == 200
    assert response.json()["kind"] == "interview_audio_analysis"


def test_learning_feed_answer_and_mistakes():
    feed = client.post(
        "/learning/feed",
        json={"user_id": "learning-user", "target_position": "Java 开发工程师", "limit": 4},
    )
    assert feed.status_code == 200
    body = feed.json()
    assert body["questions"]
    assert body["insights"]
    assert body["rag_status"]["provider"] == "local"
    assert body["rag_status"]["retrieved"] >= 0

    question_id = body["questions"][0]["id"]
    answer = client.post(
        "/learning/answers",
        json={
            "user_id": "learning-user",
            "question_id": question_id,
            "answer_text": "我会先说明背景，再讲方案和结果。",
            "answer_mode": "text",
            "target_position": "Java 开发工程师",
        },
    )
    assert answer.status_code == 200
    feedback = answer.json()["feedback"]
    assert feedback["score"] >= 0
    assert feedback["source"] in {"rules", "rules_fallback", "llm"}

    mistakes = client.get("/learning/mistakes/learning-user")
    assert mistakes.status_code == 200
    assert "mistakes" in mistakes.json()


def test_audio_analysis_fallback_keeps_warning_and_sample_without_fake_deposit(tmp_path: Path):
    audio = tmp_path / "interview.wav"
    audio.write_bytes(b"fake wav")

    with audio.open("rb") as handle:
        response = client.post(
            "/interviews/audio",
            files={"file": ("interview.wav", handle, "audio/wav")},
            data={"user_id": "audio-feed-user"},
        )

    assert response.status_code == 200
    task = client.get(f"/tasks/{response.json()['id']}").json()
    report = review_db.get_review(task["review_report_id"])
    assert report is not None
    assert report.segments[0].speaker == "system"
    assert "期待您的真实录音上传分析" in report.summary
    for leaked_word in ["预分析", "语音识别", "音频切分", "ffmpeg", "低置信"]:
        assert leaked_word not in report.summary
    assert any(segment.speaker == "interviewer" and "请介绍你最有代表性的项目" in segment.text for segment in report.segments)
    assert report.rag_diagnostics == []
    assert report.captured_questions == []

    questions = client.get("/questions").json()["questions"]
    fake_titles = [item["title"] for item in questions if item["source"] == "real_interview"]
    assert not any("请介绍你最有代表性的项目" in title for title in fake_titles)
