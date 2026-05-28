from pathlib import Path

from docx import Document

from app.agents.resume_agent import ResumeEvaluationAgent
from app.services.ai_clients import LLMCallResult


def test_ai_target_changes_resume_fit_and_recommendations(tmp_path: Path):
    resume = tmp_path / "ai_resume.docx"
    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("求职意向：AI算法工程师")
    doc.add_paragraph("技能：Python PyTorch TensorFlow SQL 机器学习 深度学习")
    doc.add_paragraph("项目经历：基于 Transformer 的文本分类项目，负责模型训练和评估。")
    doc.save(str(resume))

    report = ResumeEvaluationAgent().analyze(
        resume,
        target_position="AI算法工程师",
        user_instruction="请重点告诉我项目经历怎么优化。",
    )

    assert report.job_fit["target_position"] == "AI算法工程师"
    assert "PyTorch" in report.job_fit["matched_skills"]
    assert report.job_fit["fit_score"] > 50
    assert any("项目经历" in item or "项目" in item for item in report.recommendations)
class FakeSuccessLLM:
    def analyze_resume_report(self, profile, target_position, user_instruction=None, job_description=None):
        return LLMCallResult(
            ok=True,
            data={
                "quality_score": 91,
                "dimension_scores": {
                    "专业技能": 88,
                    "项目经验": 86,
                    "岗位匹配": 82,
                    "表达量化": 70,
                    "教育背景": 80,
                    "软技能": 76,
                },
                "job_fit": {
                    "target_position": target_position,
                    "fit_score": 83,
                    "expected_skills": ["Python", "PyTorch"],
                    "matched_skills": ["Python"],
                    "missing_skills": ["PyTorch"],
                },
                "recommendations": ["补充模型评估指标和实验对比。"],
                "jd_diagnosis": {
                    "enabled": True,
                    "match_rate": 83,
                    "core_requirements": ["Python", "PyTorch"],
                    "matched_items": ["Python"],
                    "missing_items": ["PyTorch"],
                    "suggestions": ["把项目经历改写到 JD 里的模型评估要求上。"],
                },
                "interview_risks": [
                    {
                        "risk_point": "模型评估",
                        "question": "请解释你的模型评估指标如何选择。",
                        "why_it_matters": "技术面会追问指标合理性。",
                        "defense_tip": "准备数据集、指标和对比实验。",
                        "severity": "medium",
                    }
                ],
                "logic_gaps": [
                    {
                        "issue": "技能与项目证据不足",
                        "evidence": "技能栏出现 PyTorch，项目中缺少模型细节。",
                        "suggestion": "补充 PyTorch 在项目中的具体用途。",
                        "severity": "medium",
                    }
                ],
                "reading_experience": {
                    "signal_to_noise_score": 78,
                    "cliches": [],
                    "density_notes": [],
                    "suggestions": ["保持项目符号清晰。"],
                },
                "star_optimizations": [
                    {
                        "before": "做了模型训练。",
                        "after": "建议改写为：使用 PyTorch 训练分类模型，并补充准确率、召回率等结果。",
                        "action_note": "补充行动和结果。",
                    }
                ],
            },
        )


class FakeFailingLLM:
    def analyze_resume_report(self, profile, target_position, user_instruction=None, job_description=None):
        return LLMCallResult(ok=False, error="llm http 500")


def test_resume_agent_uses_llm_multidimensional_report(tmp_path: Path):
    resume = tmp_path / "llm_resume.docx"
    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("技能：Python SQL")
    doc.add_paragraph("项目经历：参与智能诊断模型训练。")
    doc.save(str(resume))

    report = ResumeEvaluationAgent(llm=FakeSuccessLLM()).analyze(
        resume,
        target_position="AI算法工程师",
        user_instruction="重点优化项目经历",
        job_description="要求 Python、PyTorch、模型评估经验",
    )

    assert report.analysis_engine == "llm"
    assert report.quality_score == 91
    assert report.jd_diagnosis["missing_items"] == ["PyTorch"]
    assert report.interview_risks[0]["question"]
    assert report.star_optimizations[0]["after"]


def test_resume_agent_falls_back_when_llm_fails(tmp_path: Path):
    resume = tmp_path / "fallback_resume.docx"
    doc = Document()
    doc.add_paragraph("李四")
    doc.add_paragraph("技能：Python SQL")
    doc.add_paragraph("项目经历：负责数据分析。")
    doc.save(str(resume))

    report = ResumeEvaluationAgent(llm=FakeFailingLLM()).analyze(
        resume,
        target_position="数据分析师",
        user_instruction="请优化项目经历",
    )

    assert report.analysis_engine == "rules_fallback"
    assert report.llm_error == "llm http 500"
    assert report.recommendations


def test_resume_agent_ignores_isolated_headings_and_vague_words(tmp_path: Path):
    resume = tmp_path / "headings_resume.docx"
    doc = Document()
    doc.add_paragraph("王五")
    doc.add_paragraph("求职意向：后端开发工程师")
    doc.add_paragraph("技能：Java Redis MySQL Netty")
    doc.add_paragraph("项目经历")
    doc.add_paragraph("【项目背景】")
    doc.add_paragraph("优化")
    doc.add_paragraph("提升")
    doc.add_paragraph("项目名称：分布式 IM 后端系统")
    doc.add_paragraph("负责使用 Netty 和 Redis 实现消息路由，使用 MySQL 存储会话记录，并完成 10 万连接压测。")
    doc.save(str(resume))

    report = ResumeEvaluationAgent(llm=FakeFailingLLM()).analyze(
        resume,
        target_position="后端开发工程师",
        user_instruction="请优化项目经历",
    )

    risk_text = "\n".join(item["risk_point"] + item["question"] for item in report.interview_risks)
    star_before = report.star_optimizations[0]["before"]
    assert "【项目背景】" not in risk_text
    assert "你在简历中提到“优化”" not in risk_text
    assert "你在简历中提到“提升”" not in risk_text
    assert "Netty" in risk_text or "Redis" in risk_text
    assert "【项目背景】" not in star_before
    assert "分布式 IM 后端系统" in star_before or "Netty" in star_before
