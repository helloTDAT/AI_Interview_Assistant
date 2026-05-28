from __future__ import annotations

import re
from pathlib import Path

from app.core.config import settings
from app.models import ResumeFileKind, ResumeProfile
from app.services.ai_clients import OCRClient


class ResumeParser:
    """Parse resumes into a compact profile without treating section titles as evidence."""

    section_title_tokens = {
        "个人信息",
        "实习经历",
        "工作经历",
        "实践经历",
        "教育经历",
        "项目经历",
        "项目背景",
        "工作内容",
        "项目描述",
        "项目技术栈",
        "技术栈",
        "技术架构",
        "个人职责",
        "项目目标和预期",
        "项目成果",
        "专业技能",
        "个人特色",
        "竞赛获奖",
        "求职意向",
        "证书",
        "获奖",
        "荣誉",
    }
    stop_section_tokens = {"教育经历", "项目经历", "专业技能", "个人特色", "证书", "获奖", "荣誉", "求职意向"}

    def __init__(self, ocr_client: OCRClient | None = None) -> None:
        self.ocr_client = ocr_client or OCRClient()

    def parse(self, path: Path) -> ResumeProfile:
        kind = self.detect_kind(path)
        if kind == ResumeFileKind.pdf:
            raw_text, confidence, warnings = self._parse_pdf(path)
        elif kind == ResumeFileKind.docx:
            raw_text, confidence, warnings = self._parse_docx(path)
        elif kind == ResumeFileKind.image:
            result = self.ocr_client.extract_from_image(path)
            raw_text, confidence, warnings = result.text, result.confidence, []
            if result.confidence < 0.6:
                warnings.append("图片 OCR 置信度较低，需要用户或讲师确认关键字段。")
        else:
            raw_text, confidence, warnings = "", 0.0, ["不支持的简历格式。"]
        return self._build_profile(raw_text, kind, confidence, warnings)

    def detect_kind(self, path: Path) -> ResumeFileKind:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return ResumeFileKind.pdf
        if suffix == ".docx":
            return ResumeFileKind.docx
        if suffix in {".png", ".jpg", ".jpeg"}:
            return ResumeFileKind.image
        return ResumeFileKind.unsupported

    def _parse_pdf(self, path: Path) -> tuple[str, float, list[str]]:
        warnings: list[str] = []
        text = ""
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as exc:  # pragma: no cover - depends on external files
            warnings.append(f"PDF 文本抽取失败：{exc}")
        if len(text) < settings.min_pdf_text_chars:
            ocr = self.ocr_client.extract_from_pdf_pages(path)
            warnings.append("PDF 文本层过少，已切换到扫描版 OCR 解析路径。")
            if ocr.confidence < 0.6:
                warnings.append("PDF OCR 置信度较低，需要人工确认。")
            return ocr.text, ocr.confidence, warnings
        return text, 0.86, warnings

    def _parse_docx(self, path: Path) -> tuple[str, float, list[str]]:
        try:
            from docx import Document

            doc = Document(str(path))
            chunks: list[str] = []
            chunks.extend(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        chunks.append(row_text)
            return "\n".join(chunks).strip(), 0.9, []
        except Exception as exc:  # pragma: no cover - depends on external files
            return "", 0.0, [f"DOCX 解析失败：{exc}"]

    def _build_profile(
        self,
        raw_text: str,
        kind: ResumeFileKind,
        confidence: float,
        warnings: list[str],
    ) -> ResumeProfile:
        lines = [self._normalize_line(line) for line in raw_text.splitlines() if line.strip()]
        name = self._guess_name(lines)
        contact = self._find_contact(raw_text)
        skills = self._extract_skills(raw_text)
        return ResumeProfile(
            candidate_name=name,
            contact=contact,
            education=self._extract_section(lines, ["教育", "教育经历", "学历"]),
            projects=self._extract_projects(lines),
            internships=self._extract_internships(lines),
            skills=skills,
            certificates=self._extract_section(lines, ["证书", "获奖", "荣誉"]),
            target_position=self._guess_target(raw_text, skills),
            raw_text=raw_text,
            parse_confidence=confidence,
            source_kind=kind,
            warnings=warnings,
        )

    def _guess_name(self, lines: list[str]) -> str | None:
        for line in lines[:5]:
            clean = re.sub(r"\s+", "", line)
            if 2 <= len(clean) <= 8 and not re.search(r"简历|电话|邮箱|求职|学历|个人信息", clean):
                return clean
        return None

    def _find_contact(self, text: str) -> str | None:
        email = re.search(r"[\w.\-+]+@[\w.\-]+\.\w+", text)
        phone = re.search(r"1[3-9]\d{9}", text)
        contacts = [m.group(0) for m in [phone, email] if m]
        return " / ".join(contacts) if contacts else None

    def _extract_skills(self, text: str) -> list[str]:
        known = [
            "Python",
            "Java",
            "C++",
            "SQL",
            "MySQL",
            "Redis",
            "Docker",
            "Linux",
            "React",
            "Vue",
            "FastAPI",
            "Spring",
            "SpringCloud",
            "Netty",
            "RocketMQ",
            "Zookeeper",
            "Nacos",
            "Dubbo",
            "PyTorch",
            "TensorFlow",
            "NLP",
            "Transformer",
            "OpenCV",
            "LangChain",
            "WebRTC",
            "asyncio",
            "Whisper",
            "算法",
            "数据分析",
            "机器学习",
            "深度学习",
        ]
        lowered = text.lower()
        return [skill for skill in known if skill.lower() in lowered]

    def _extract_projects(self, lines: list[str]) -> list[str]:
        blocks: list[str] = []
        index = 0
        while index < len(lines):
            line = lines[index]
            if not self._is_project_start(line):
                index += 1
                continue

            block: list[str] = []
            while index < len(lines):
                current = lines[index]
                if block and self._is_project_start(current):
                    break
                if block and self._is_stop_section(current, extra={"实习经历", "工作经历", "实践经历"}):
                    break
                if not self._is_title_only(current):
                    block.append(current)
                index += 1

            compact = self._compact_block(block)
            if compact and self._has_experience_signal(compact):
                blocks.append(compact)
            continue

        return list(dict.fromkeys(blocks))[:6]

    def _extract_internships(self, lines: list[str]) -> list[str]:
        blocks: list[str] = []
        index = 0
        while index < len(lines):
            if not self._normalized_title(lines[index]).startswith(("实习经历", "工作经历", "实践经历")):
                index += 1
                continue

            index += 1
            block: list[str] = []
            while index < len(lines):
                current = lines[index]
                if self._is_stop_section(current, extra={"教育经历", "项目经历", "专业技能", "个人特色"}):
                    break
                if not self._is_title_only(current):
                    block.append(current)
                index += 1

            compact = self._compact_block(block)
            if compact and self._has_experience_signal(compact):
                blocks.append(compact)
            continue

        return list(dict.fromkeys(blocks))[:4]

    def _extract_section(self, lines: list[str], keywords: list[str]) -> list[str]:
        matched: list[str] = []
        for index, line in enumerate(lines):
            normalized = self._normalized_title(line)
            if any(keyword in normalized for keyword in keywords):
                for candidate in lines[index : min(index + 4, len(lines))]:
                    if not self._is_title_only(candidate):
                        matched.append(candidate)
        return list(dict.fromkeys(matched))[:6]

    def _guess_target(self, text: str, skills: list[str]) -> str | None:
        match = re.search(r"求职意向[:：\s]*(.+)", text)
        if match:
            return match.group(1).strip()[:40]
        if any(skill in skills for skill in ["Python", "Java", "MySQL", "Redis", "FastAPI", "Spring", "Netty"]):
            return "后端开发工程师"
        if any(skill in skills for skill in ["React", "Vue"]):
            return "前端开发工程师"
        return None

    def _normalize_line(self, line: str) -> str:
        return re.sub(r"\s+", " ", line.strip())

    def _normalized_title(self, line: str) -> str:
        text = re.sub(r"^[\s•·●■\-]+", "", line.strip())
        text = text.strip("【】[]（）()：: ")
        text = re.sub(r"\s+", "", text)
        return text

    def _is_title_only(self, line: str) -> bool:
        normalized = self._normalized_title(line)
        if normalized in self.section_title_tokens:
            return True
        return bool(re.fullmatch(r"项目经历\d*", normalized))

    def _is_project_start(self, line: str) -> bool:
        normalized = self._normalized_title(line)
        if re.match(r"项目经历\d*", normalized):
            return True
        if normalized.startswith("项目名称"):
            return True
        return False

    def _is_stop_section(self, line: str, extra: set[str] | None = None) -> bool:
        normalized = self._normalized_title(line)
        tokens = self.stop_section_tokens | (extra or set())
        if normalized in tokens:
            return True
        return bool(re.fullmatch(r"(教育经历|专业技能|个人特色|证书|荣誉|获奖)", normalized))

    def _compact_block(self, lines: list[str]) -> str:
        cleaned = [line for line in lines if line and not self._is_title_only(line)]
        return "\n".join(cleaned).strip()[:900]

    def _has_experience_signal(self, text: str) -> bool:
        has_action = bool(re.search(r"负责|参与|设计|实现|完成|主导|优化|搭建|开发|重构|接入|迁移|压测|部署", text))
        has_anchor = bool(
            re.search(
                r"\d|%|Redis|MySQL|Netty|KMS|RocketMQ|WebRTC|Python|Java|SQL|Docker|Spring|"
                r"平台|系统|模型|模块|接口|服务|数据库|消息|缓存|前端|后端|业务|用户|订单",
                text,
                re.I,
            )
        )
        return has_action and has_anchor and len(text.strip()) >= 18
